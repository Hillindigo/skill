import argparse
import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def q(tag):
    return qn('w:' + tag)


def set_font(run, size=10.5, bold=False, color='333333'):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(q('eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_margins(cell, top=80, bottom=80, start=120, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for name, value in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        node = tcMar.find(q(name))
        if node is None:
            node = OxmlElement('w:' + name)
            tcMar.append(node)
        node.set(q('w'), str(value))
        node.set(q('type'), 'dxa')


def set_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(q('tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(q('w'), str(width))
    tcW.set(q('type'), 'dxa')


def set_geometry(table, widths):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(q('tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(q('w'), str(sum(widths)))
    tblW.set(q('type'), 'dxa')
    tblInd = tblPr.find(q('tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd')
        tblPr.append(tblInd)
    tblInd.set(q('w'), '120')
    tblInd.set(q('type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(q('w'), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_width(cell, widths[index])
            set_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header(row):
    trPr = row._tr.get_or_add_trPr()
    node = OxmlElement('w:tblHeader')
    node.set(q('val'), 'true')
    trPr.append(node)


def fmt_time(value):
    value = float(value)
    minutes = int(value // 60)
    seconds = value % 60
    return f'{minutes:02d}:{seconds:05.2f}'


def load_data(path):
    with open(path, encoding='utf-8-sig') as handle:
        payload = json.load(handle)
    if isinstance(payload, list):
        return {'segments': payload}
    return payload


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_font(r, 20, True, '1F4D78')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(subtitle)
    set_font(r, 12, False, '666666')


def add_meta(doc, data, mode, source):
    meta = [
        ('版本', mode),
        ('创作者', data.get('author') or data.get('creator') or '未提供'),
        ('原始链接', source or data.get('source_url') or '未提供'),
        ('转写方式', '本地 faster-whisper 中文语音转写'),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + '：')
        set_font(r, 10, True, '1F4D78')
        r = p.add_run(str(value))
        set_font(r, 10, False, '666666')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('说明：')
    set_font(r, 10, True, '9C6500')
    r = p.add_run('以下内容来自视频语音转写，可能存在少量识别误差，不等同于平台官方字幕。')
    set_font(r, 10, False, '666666')


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(section, attr, Inches(1))
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal._element.rPr.rFonts.set(q('eastAsia'), 'Microsoft YaHei')
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ('Heading 1', 16, '2E74B5', 18, 10),
        ('Heading 2', 13, '2E74B5', 12, 6),
    ]:
        style = doc.styles[name]
        style.font.name = 'Calibri'
        style._element.rPr.rFonts.set(q('eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    return doc


def build_full(data, output, source):
    doc = setup_doc()
    add_title(doc, data.get('title') or '抖音视频完整字幕', '完整逐句字幕版')
    add_meta(doc, data, '完整逐句字幕版', source)
    doc.add_paragraph('完整字幕', style='Heading 1')
    segments = data.get('segments', [])
    table = doc.add_table(rows=len(segments) + 1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_geometry(table, [1500, 7860])
    mark_header(table.rows[0])
    for i, label in enumerate(('时间', '字幕')):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(label)
        set_font(r, 10.5, True, '1F4D78')
    for row, segment in zip(table.rows[1:], segments):
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(fmt_time(segment.get('start', 0)) + '\n' + fmt_time(segment.get('end', 0)))
        set_font(r, 9.2, False, '666666')
        p = row.cells[1].paragraphs[0]
        r = p.add_run(segment.get('text', '').strip())
        set_font(r, 10.2)
    doc.save(output)


def build_original(data, output, source, sections_path):
    doc = setup_doc()
    add_title(doc, data.get('title') or '抖音视频字幕', '博主原话文章版 · 分点整理')
    add_meta(doc, data, '博主原话文章版', source)
    with open(sections_path, encoding='utf-8-sig') as handle:
        sections = json.load(handle)
    segments = data.get('segments', [])
    doc.add_paragraph('原话整理', style='Heading 1')
    note = doc.add_paragraph('正文仅按主题分段、添加标点和排版；标题为整理用，不代表新增观点。')
    for run in note.runs:
        set_font(run, 10, False, '666666')
    for section in sections:
        doc.add_paragraph(section['heading'], style='Heading 2')
        start = int(section['start'])
        end = int(section['end'])
        text = [str(item.get('text', '')).strip() for item in segments[start:end] if str(item.get('text', '')).strip()]
        group = int(section.get('group', 8))
        for offset in range(0, len(text), group):
            paragraph = doc.add_paragraph('，'.join(text[offset:offset + group]).strip('，') + '。')
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.line_spacing = 1.28
            for run in paragraph.runs:
                set_font(run, 11, False, '333333')
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', required=True, help='transcript JSON')
    parser.add_argument('--output', required=True, help='output DOCX')
    parser.add_argument('--mode', choices=['full', 'original-article'], default='full')
    parser.add_argument('--source', default='')
    parser.add_argument('--sections', help='JSON list with heading/start/end for original-article')
    args = parser.parse_args()
    data = load_data(args.input)
    os.makedirs(os.path.dirname(os.path.abspath(args.output)), exist_ok=True)
    if args.mode == 'full':
        build_full(data, args.output, args.source)
    else:
        if not args.sections:
            parser.error('--sections is required for original-article')
        build_original(data, args.output, args.source, args.sections)
    print(args.output)


if __name__ == '__main__':
    main()
